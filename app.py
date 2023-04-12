import base64
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from doc_reader import read_questions_from_docx
from werkzeug.utils import secure_filename
import os
from io import BytesIO
from docx import Document
import io

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = "your_secret_key"

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

questions = []


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files["file"]

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(file_path)

            new_questions = read_questions_from_docx(file_path)
            os.remove(file_path)

            global questions
            questions.extend(new_questions)

            return render_template("index.html", questions=questions)

    return render_template("index.html", questions=questions)


@app.route('/upload_file', methods=['POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)

        file = request.files['file']
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            new_questions = read_questions_from_docx(os.path.join(app.config['UPLOAD_FOLDER'], filename))

            # Add print statements to check the parsed questions
            print("Uploaded questions:")
            for question in new_questions:
                print(question)

            global questions
            questions.extend(new_questions)

            return redirect(url_for('index'))
        else:
            flash('Allowed file types are docx')
            return redirect(request.url)

    return redirect(url_for('index'))



@app.route('/edit_questions', methods=['POST'])
def edit_questions():
    for key, value in request.form.items():
        if key.startswith('question_type_'):
            index = int(key.split('_')[-1])
            question_type = value
            question_index = int(request.form[f'question_index_{index}'])
            question_text = request.form[f'question_text_{index}']
            question_difficulty = int(request.form[f'question_difficulty_{index}'])
            questions[question_index]['type'] = question_type
            questions[question_index]['text'] = question_text
            questions[question_index]['difficulty'] = question_difficulty

    return redirect(url_for('index'))


@app.route('/generate_tickets', methods=['GET', 'POST'])
def generate_tickets():
    num_tickets = int(request.form.get('num_tickets', 0))
    target_difficulty = request.form.get('target_difficulty', '')

    if target_difficulty == '':
        # Calculate the average difficulty if target_difficulty is not provided
        target_difficulty = None
    else:
        target_difficulty = int(target_difficulty)

    print(f"Target difficulty: {target_difficulty}")
    print("Questions Before generation", questions)
    generated_tickets = generate_balanced_tickets(questions, num_tickets, target_difficulty)

    # Add a unique ID to each ticket
    tickets = []
    for i, ticket in enumerate(generated_tickets):
        tickets.append({"id": f"{i}", "content": ticket})

    return render_template('tickets.html', tickets=tickets)


def calculate_difficulty_distance(ticket, target_difficulty):
    return abs(sum(q['difficulty'] for q in ticket) - target_difficulty)


def count_overlapping_questions(ticket1, ticket2):
    return len(set(q['text'] for q in ticket1) & set(q['text'] for q in ticket2))
def get_combinations(theoretical_questions, practical_questions, target_difficulty, k=3, depth=0, path=None):
    if path is None:
        path = []

    if depth == k:
        print(f"Depth reached: {depth}, current path: {path}")
        for practical_question in practical_questions:
            ticket = path + [practical_question]
            ticket_difficulty = sum(q['difficulty'] for q in ticket)
            print(f"Trying ticket: {ticket}, difficulty: {ticket_difficulty}")
            if abs(ticket_difficulty - target_difficulty) <= k:
                print(f"Yielding ticket: {ticket}")
                yield ticket
    else:
        for i, question in enumerate(theoretical_questions):
            new_path = path + [question]
            remaining_questions = theoretical_questions[i + 1:]
            yield from get_combinations(remaining_questions, practical_questions, target_difficulty, k, depth + 1,
                                         new_path)

def generate_balanced_tickets(questions, num_tickets, target_difficulty):
    theoretical_questions = [q for q in questions if q['type'] == 'theoretical']
    practical_questions = [q for q in questions if q['type'] == 'practical']

    print(f"Theoretical questions: {theoretical_questions}")
    print(f"Practical questions: {practical_questions}")

    all_combinations = []

    if target_difficulty is None:
        target_difficulty = int(sum(q['difficulty'] for q in questions))

    print(f"Target difficulty: {target_difficulty}")


    for practical_question in practical_questions:
        all_combinations.extend(list(get_combinations(theoretical_questions, [practical_question], target_difficulty)))

    print(f"All ticket combinations: {all_combinations}")


    generated_tickets = []
    while num_tickets > 0 and all_combinations:
        min_diff = float('inf')
        best_ticket = None
        for ticket in all_combinations:
            diff = abs(sum(q['difficulty'] for q in ticket) - target_difficulty)
            if diff < min_diff:
                min_diff = diff
                best_ticket = ticket

        if best_ticket:
            generated_tickets.append(best_ticket)
            all_combinations.remove(best_ticket)
            num_tickets -= 1
        else:
            break

    print(f"Generated {len(generated_tickets)} ticket(s)")

    return generated_tickets




@app.route('/save_tickets', methods=['POST'])
def save_tickets():
    ticket_data = request.json['ticketsData']

    # Create a new Word document
    doc = Document()

    # Iterate through the ticket data and add the questions to the Word document
    for index, ticket in enumerate(ticket_data):
        doc.add_heading(f'Ticket {index + 1}', level=2)
        for q_index, question in enumerate(ticket):
            doc.add_paragraph(f'{q_index + 1}. {question["text"]} (Difficulty: {question["difficulty"]})')

        # Add a page break after each ticket, except the last one
        if index < len(ticket_data) - 1:
            doc.add_page_break()

    # Save the Word document to a file
    doc.save('tickets.docx')

    # Send the saved file to the client
    return send_file('tickets.docx', as_attachment=True, download_name='tickets.docx')


if __name__ == '__main__':
    app.run(debug=True)
